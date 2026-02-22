from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document


CATEGORY_TITLES = {
    "missing_items": "必填项缺失",
    "risky_phrases": "风险用词",
    "closure_issues": "闭环问题",
    "logic_warnings": "逻辑/一致性提醒",
}

SEV_LABEL = {
    "high": "高",
    "medium": "中",
    "low": "低",
}


def build_docx_report(report: dict[str, Any], doc_type: str, source: str = "manual") -> bytes:
    doc = Document()

    doc.add_heading("监理文书合规校验报告", level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"文书类型：{doc_type}")
    doc.add_paragraph(f"文本来源：{source}")

    summary = report.get("summary", {})
    doc.add_heading("一、汇总", level=2)
    doc.add_paragraph(
        f"总提示：{summary.get('total', 0)}；高：{summary.get('high', 0)}；"
        f"中：{summary.get('medium', 0)}；低：{summary.get('low', 0)}"
    )

    findings = report.get("findings", [])
    doc.add_heading("二、问题明细", level=2)
    if not findings:
        doc.add_paragraph("未发现明显问题（建议人工复核）。")
    else:
        grouped: dict[str, list[dict[str, Any]]] = {}
        for item in findings:
            grouped.setdefault(str(item.get("category", "unknown")), []).append(item)

        for key in ("missing_items", "risky_phrases", "closure_issues", "logic_warnings"):
            items = grouped.get(key, [])
            if not items:
                continue
            title = CATEGORY_TITLES.get(key, key)
            doc.add_heading(f"{title}（{len(items)}）", level=3)
            for i, f in enumerate(items, 1):
                sev = SEV_LABEL.get(str(f.get("severity", "low")), str(f.get("severity", "low")))
                doc.add_paragraph(f"{i}. [{sev}] {f.get('title', '')}")
                if f.get("quote"):
                    doc.add_paragraph(f"触发：{f['quote']}")
                if f.get("reason"):
                    doc.add_paragraph(f"原因：{f['reason']}")
                if f.get("suggestion"):
                    doc.add_paragraph(f"建议：{f['suggestion']}")

    suggestions = report.get("copy_ready_suggestions", [])
    doc.add_heading("三、可复制建议句", level=2)
    if not suggestions:
        doc.add_paragraph("无")
    else:
        for i, s in enumerate(suggestions, 1):
            doc.add_paragraph(f"{i}. {s}")

    rewrite = str(report.get("full_text_rewrite", "") or "").strip()
    if rewrite:
        doc.add_heading("四、建议替换后的完整日志（草案）", level=2)
        doc.add_paragraph(rewrite)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
